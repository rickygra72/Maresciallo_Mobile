import streamlit as st
import os
import re
from google import genai
from google.genai import types

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import reportlab
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# Configurazione Pagina Streamlit
st.set_page_config(
    page_title="Maresciallo AI - Unità Investigativa",
    page_icon="🛡️",
    layout="centered",
    initial_sidebar_state="expanded"
)

PASSWORD_APPLICATIVO = "GdiF_117"
DEFAULT_COST_INPUT_1M_EUR = 0.069
DEFAULT_COST_OUTPUT_1M_EUR = 0.276

# Gestione Autenticazione con campo password
if "autenticato" not in st.session_state:
    st.session_state.autenticato = False

if not st.session_state.autenticato:
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; color: #FFD700;'>🔒 SISTEMA CRITTOGRAFATO - ACCESSO RISERVATO</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #8b949e;'>Inserire la password di sblocco per accedere alla plancia operativa:</p>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        pwd_input = st.text_input("Password:", type="password", label_visibility="collapsed", placeholder="Inserisci password...")
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🔓 SBLOCCA TERMINALE"):
            if pwd_input.strip() == PASSWORD_APPLICATIVO:
                st.session_state.autenticato = True
                st.rerun()
            else:
                st.error("❌ Password errata.")
    st.stop()

# Inizializzazione Stati di Sessione
if "storico_chat" not in st.session_state:
    st.session_state.storico_chat = []
if "sess_token_in" not in st.session_state:
    st.session_state.sess_token_in = 0
if "sess_token_out" not in st.session_state:
    st.session_state.sess_token_out = 0
if "sess_costo_eur" not in st.session_state:
    st.session_state.sess_costo_eur = 0.0

# --- SIDEBAR OPERATIVA ---
with st.sidebar:
    st.markdown("### 🛡️ Maresciallo AI [Mobile]")
    
    modalita = st.selectbox(
        "🏢 Reparto Operativo:",
        options=["PEF", "PG"],
        format_func=lambda x: "🟡 Polizia Economico-Finanziaria (P.E.F.)" if x == "PEF" else "🔵 Polizia Giudiziaria (P.G.)"
    )
    
    api_key_input = st.text_input("Chiave API Google:", type="password", value=os.environ.get("GEMINI_API_KEY", ""))
    
    stile_interfaccia = st.selectbox(
        "💬 Stile Interfaccia:",
        options=["Collega (Naturale & Dettagliato)", "Sintetico (Solo Dati Essenziali)"]
    )
    
    st.markdown("---")
    st.markdown("### 📂 ALLEGATI")
    uploaded_files = st.file_uploader(
        "Carica atti, verbali, estratti conto o file contabili",
        type=["pdf", "xlsx", "txt", "png", "jpg", "jpeg", "docx"],
        accept_multiple_files=True
    )
    
    presenza_immagini = False
    if uploaded_files:
        for f in uploaded_files:
            if f.name.split('.')[-1].lower() in ["jpg", "jpeg", "png"]:
                presenza_immagini = True
        if presenza_immagini:
            st.warning("⚠️ Rilevate immagini/scansioni: l'elaborazione visiva (OCR) inciderà sui token.")

    st.markdown("---")
    st.markdown("### 🖨️ ESPORTAZIONE ATTI")
    
    # Sezione Download sempre attiva se c'è almeno un messaggio nello storico
    if len(st.session_state.storico_chat) > 0:
        # Generazione Word in memoria per il download
        if Document:
            doc = Document()
            doc.styles['Normal'].font.name = 'Arial'
            titolo_rep = "RELAZIONE DI ANALISI ISPETTIVA (P.E.F.)" if modalita == "PEF" else "ANNOTAZIONE DI POLIZIA GIUDIZIARIA (P.G.)"
            doc.add_heading(titolo_rep, level=1)
            for m in st.session_state.storico_chat:
                if m['role'] == 'assistant':
                    doc.add_paragraph(m['content'])
                    doc.add_paragraph("-" * 40)
            
            doc_path = "relazione_maresciallo.docx"
            doc.save(doc_path)
            with open(doc_path, "rb") as f_docx:
                st.download_button("📥 Scarica Word (.docx)", f_docx, file_name="Relazione_Maresciallo.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Generazione PDF in memoria per il download
        if HAS_REPORTLAB:
            pdf_path = "relazione_maresciallo.pdf"
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter, rightMargin=20, leftMargin=20, topMargin=30, bottomMargin=30)
            styles = getSampleStyleSheet()
            style_testo = ParagraphStyle('TestoReport', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=10)
            
            story = [Paragraph("RELAZIONE TECNICA ISPETTIVA", styles['Heading1']), Spacer(1, 8)]
            for m in st.session_state.storico_chat:
                if m['role'] == 'assistant':
                    for riga in m['content'].split('\n'):
                        if riga.strip():
                            story.append(Paragraph(riga.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), style_testo))
                    story.append(Spacer(1, 10))
            pdf_doc.build(story)
            
            with open(pdf_path, "rb") as f_pdf:
                st.download_button("📥 Scarica PDF Formale", f_pdf, file_name="Relazione_Maresciallo.pdf", mime="application/pdf")
    else:
        st.caption("I pulsanti di download compariranno qui non appena l'IA avrà prodotto la prima risposta.")

    st.markdown("---")
    st.markdown("### 💶 CONSUMI SESSIONE")
    tot_token = st.session_state.sess_token_in + st.session_state.sess_token_out
    st.info(f"• Token Elaborati: {tot_token:,}\n• Spesa Stimata: {st.session_state.sess_costo_eur:.4f} €")
    
    st.markdown("---")
    if st.button("🔄 Nuova Indagine (Reset)"):
        st.session_state.storico_chat = []
        st.session_state.sess_token_in = 0
        st.session_state.sess_token_out = 0
        st.session_state.sess_costo_eur = 0.0
        st.rerun()

# --- CORPO PRINCIPALE CHAT ---
titolo_plancia = "Polizia Economico-Finanziaria (P.E.F.)" if modalita == "PEF" else "Polizia Giudiziaria (P.G.)"
st.markdown(f"## 🛡️ Unità Investigativa - {titolo_plancia}")
st.markdown("<div style='font-size: 0.9em; color: #8b949e; margin-bottom: 15px;'>Protocollo di analisi forense, riscontro contabile e accertamenti bancari attivo.</div>", unsafe_allow_html=True)

# Visualizzazione nativa e pulita dello storico chat (usa i componenti standard di Streamlit per la massima leggibilità)
for msg in st.session_state.storico_chat:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Input Utente
prompt_utente = st.chat_input("Inserisci il quesito o l'atto da esaminare...")

if prompt_utente:
    if not api_key_input:
        st.error("⚠️ Inserisci la Chiave API Google nella barra laterale per procedere.")
        st.stop()
        
    st.session_state.storico_chat.append({"role": "user", "content": prompt_utente})
    with st.chat_message("user"):
        st.markdown(prompt_utente)
        
    with st.chat_message("assistant"):
        with st.spinner("Elaborazione e riscontro in corso..."):
            try:
                testo_atti = ""
                immagini = []
                
                if uploaded_files:
                    for f in uploaded_files:
                        ext = f.name.split('.')[-1].lower()
                        if ext in ["jpg", "jpeg", "png"]:
                            mime = "image/jpeg" if ext in ["jpg", "jpeg"] else "image/png"
                            immagini.append(types.Part.from_bytes(data=f.read(), mime_type=mime))
                            continue
                        
                        testo_atti += f"\n[INIZIO FILE: {f.name}]\n"
                        try:
                            if ext == "txt":
                                testo_atti += f.read().decode("utf-8", errors="ignore")
                            elif ext == "pdf" and pypdf:
                                reader = pypdf.PdfReader(f)
                                for num_p, page in enumerate(reader.pages, start=1):
                                    t = page.extract_text()
                                    if t: testo_atti += f"--- Pagina {num_p} ---\n" + t + "\n"
                            elif ext == "xlsx" and openpyxl:
                                wb = openpyxl.load_workbook(f, data_only=True)
                                for sheet_name in wb.sheetnames:
                                    ws = wb[sheet_name]
                                    testo_atti += f"[Foglio Excel: {sheet_name}]\n"
                                    for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                                        if any(row):
                                            riga_formattata = " | ".join([str(c) if c is not None else "" for c in row])
                                            testo_atti += f"Riga {row_idx}: {riga_formattata}\n"
                            elif ext == "docx" and Document:
                                doc = Document(f)
                                for p in doc.paragraphs:
                                    if p.text: testo_atti += p.text + "\n"
                        except Exception as e:
                            testo_atti += f"[Errore lettura file: {str(e)}]\n"
                        testo_atti += f"\n[FINE FILE: {f.name}]\n"

                p_completo = prompt_utente
                if testo_atti:
                    p_completo += f"\n\n[DATI ESTRATTI DAI DOCUMENTI ALLEGATI (DA LEGGERE RIGOROSAMENTE SECONDO LE INTESTAZIONI LETTERALI)]\n{testo_atti}"

                client = genai.Client(api_key=api_key_input)
                
                istruzione_stile = "Esponi con tono naturale, collaborativo e articolato, mantenendo elevata precisione tecnica."
                if "Sintetico" in stile_interfaccia:
                    istruzione_stile = "Esponi in modo estremamente sintetico, schematico e diretto, fornendo solo i dati e le conclusioni essenziali."

                if modalita == "PEF":
                    sys_prompt = (
                        "Sei un Assistente Virtuale Senior specializzato in Polizia Economico-Finanziaria (P.E.F.), Audit Forense, Diritto Tributario e Accertamenti Bancari.\n\n"
                        f"--- STILE DI INTERAZIONE ---\n{istruzione_stile}\n\n"
                        "--- PROTOCOLLO GEOMETRICO E TASSATIVO CONTRO L'INVERSIONE DEI DATI ---\n"
                        "1. VINCOLO ASSOLUTO DI LETTURA LETTERALE:\n"
                        "   • È SEVERAMENTE VIETATO invertire le operazioni ATTIVE (IVA a debito / cessioni / imponibile attivo / accrediti / versamenti) con le operazioni PASSIVE (IVA a credito / acquisti / imponibile passivo / prelevamenti).\n"
                        "   • Devi leggere esclusivamente la stringa di testo o la riga della tabella associata all'etichetta letterale originale. La parola scritta e l'etichetta di colonna hanno priorità matematica assoluta.\n\n"
                        "--- DIRETTIVE SPECIALIZZATE PER ACCERTAMENTI BANCARI E FINANZIARI (art. 32 D.P.R. 600/73 & art. 51 D.P.R. 633/72) ---\n"
                        "1. ANALISI MOVIMENTI BANCARI ED ESTRATTI CONTO:\n"
                        "   • VERSAMENTI (Accrediti): Vanno considerati, per presunzione legale (inversione dell'onere della prova), come compensi o ricavi non dichiarati, salvo rigorosa prova contraria fornita dal contribuente (es. giroconti, finanziamenti soci, liberalità, disinvestimenti giustificati).\n"
                        "   • PRELEVAMENTI (Addebiti): Verifica la qualifica soggettiva del contribuente (imprenditore vs lavoratore autonomo / professionista). Ricorda che per i professionisti/lavoratori autonomi i prelevamenti ingiustificati non costituiscono compensi occulti (Corte Cost. n. 228/2014), mentre per gli imprenditori commerciali permane la presunzione limitata alle soglie normative vigenti.\n"
                        "2. QUADRATURA E TABELLE DI RISCONTRO:\n"
                        "   • Genera tabelle di scomposizione analitica indicando chiaramente: Data, Causale, Importo, Natura (Versamento vs Prelevamento), Eventuale Giustificazione documentale e Rischio Fiscale.\n\n"
                        "--- STRUTTURA DELLA RISPOSTA ---\n"
                        "- 📌 PERIMETRO DI ANALISI ED ESAME DOCUMENTALE\n"
                        "- 📍 QUESITO / FATTISPECIE ISPETTIVA O BANCARIA\n"
                        "- 📊 TABELLA DI QUADRATURA E RISCONTRO ANALITICO (ATTIVI / PASSIVI O VERSAMENTI / PRELEVAMENTI)\n"
                        "- 🔍 RILIEVI FISCALI, PRESUNZIONI LEGALI ED EVENTUALI PROFILI PENALI (D.Lgs. 74/2000)\n"
                        "- 📝 CONCLUSIONI TECNICHE ED APPROFONDIMENTI ISPETTIVI SUGGERITI"
                    )
                else:
                    sys_prompt = (
                        "Sei un Assistente Virtuale Senior specializzato in Polizia Giudiziaria (P.G.), Diritto Penale e Procedura Penale (c.p.p.).\n\n"
                        f"--- STILE DI INTERAZIONE ---\n{istruzione_stile}\n\n"
                        "--- DIRETTIVE OPERATIVE POLIZIA GIUDIZIARIA ---\n"
                        "1. DICHIARAZIONE PERIMETRO: Inizia dichiarando la quantità esatta di atti penali ed elementi di prova esaminati.\n"
                        "2. QUALIFICAZIONE PENALE E NORMATIVA PROCEDURALE:\n"
                        "   • Inquadramento Penale Rigoroso: Qualifica giuridicamente le condotte penalmente rilevanti (D.Lgs. 74/2000, Reati Societari ex art. 2621 e ss. c.c., Reati Fallimentari/Bancarotta ex CCII, Reati contro la P.A. o il Patrimonio).\n"
                        "   • Elemento Soggettivo e Materiale: Analizza il dolo specifico di evasione o di profitto e la condotta materiale con riferimenti alle più recenti sentenze di Cassazione Penale.\n"
                        "   • Procedura Penale (c.p.p.): Verifica la regolarità degli atti di P.G. (art. 347 c.p.p. Informativa di Reato, art. 352 c.p.p. Perquisizione, art. 354 c.p.p. Sequestro probatorio/preventivo) e il rispetto tassativo delle GARANZIE DIFENSIVE ex art. 356 c.p.p. per evitare eccezioni di inutilizzabilità.\n"
                        "3. REDAZIONE ATTI DI P.G.: Fornisci indicazioni formali chiare e bozze di clausole giuridiche per la redazione di Annotazioni e Informative di Reato a perfetta tenuta dibattimentale.\n\n"
                        "--- STRUTTURA DELLA RISPOSTA ---\n"
                        "- 📌 PERIMETRO DI ANALISI ED ESAME ATTI DI P.G.\n"
                        "- 📍 REATO / FATTISPECIE PENALE CONTESTATA\n"
                        "- ⚖️ QUALIFICAZIONE GIURIDICA (Fattispecie, Dolo, Giurisprudenza & Soglie Penali)\n"
                        "- 🔍 RISCONTRO PROCEDURALE PENALE (c.p.p., Garanzie & Tenuta Probatoria)\n"
                        "- 📝 DIRETTIVE FORMALI PER LA REDAZIONE DEGLI ATTI DI P.G."
                    )

                contents = []
                for m in st.session_state.storico_chat[:-1]:
                    contents.append(types.Content(role="user" if m["role"] == "user" else "model", parts=[types.Part.from_text(text=m["content"])]))
                contents.append(types.Content(role="user", parts=[types.Part.from_text(text=p_completo)] + immagini))

                res = client.models.generate_content(
                    model='gemini-3.6-flash',
                    contents=contents,
                    config=types.GenerateContentConfig(system_instruction=sys_prompt, temperature=0.0)
                )

                costo_atto = 0.0
                if hasattr(res, 'usage_metadata') and res.usage_metadata:
                    t_in = res.usage_metadata.prompt_token_count or 0
                    t_out = res.usage_metadata.candidates_token_count or 0
                    st.session_state.sess_token_in += t_in
                    st.session_state.sess_token_out += t_out
                    costo_atto = (t_in / 1_000_000.0) * DEFAULT_COST_INPUT_1M_EUR + (t_out / 1_000_000.0) * DEFAULT_COST_OUTPUT_1M_EUR
                    st.session_state.sess_costo_eur += costo_atto

                risposta_ia = res.text + f"\n\n[💶 Spesa API per questo atto: {costo_atto:.5f} €]"
                
                st.markdown(risposta_ia)
                st.session_state.storico_chat.append({"role": "assistant", "content": risposta_ia})
                st.rerun()

            except Exception as e:
                st.error(f"⚠️ Errore durante l'elaborazione: {str(e)}")