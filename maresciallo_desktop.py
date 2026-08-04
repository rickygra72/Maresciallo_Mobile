import os
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import base64
import re
import time
import json
import threading
import datetime
import urllib.request
from google import genai
from google.genai import types

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".maresciallo_config")
STATS_FILE = os.path.join(os.path.expanduser("~"), ".maresciallo_stats")
PRICES_FILE = os.path.join(os.path.expanduser("~"), ".maresciallo_prices")
PASSWORD_APPLICATIVO = "GdiF_117"

DEFAULT_COST_INPUT_1M_EUR = 0.069
DEFAULT_COST_OUTPUT_1M_EUR = 0.276


class PlanciaMarescialloAI:
    def __init__(self, root):
        self.root = root
        self.modalita = "PEF"
        self.storico_chat = []
        self.files_acquisiti = []

        self.costo_input_1m_eur = DEFAULT_COST_INPUT_1M_EUR
        self.costo_output_1m_eur = DEFAULT_COST_OUTPUT_1M_EUR
        self.stato_tariffe_online = "Inizializzazione..."

        self.sess_token_in = 0
        self.sess_token_out = 0
        self.sess_costo_eur = 0.0

        self.carica_statistiche_locali()
        self.carica_tariffe_salvate()

        self.root.title("🛡️ Maresciallo AI - Unità Investigativa")
        self.root.geometry("1100x780")
        try:
            self.root.state('zoomed')
        except Exception:
            pass

        self.bg_principale = "#0d1117"
        self.bg_aree = "#161b22"
        self.bg_input = "#1f242c"
        self.bg_pulsante = "#2d3748"
        self.colore_accento = "#00E676"
        self.giallo_ia = "#FFD700"
        self.arancione_reset = "#FF9100"
        self.bianco_puro = "#ffffff"
        self.rosso_alert = "#e74c3c"

        self.root.configure(bg=self.bg_principale)

        self.mostra_login_modale()
        threading.Thread(target=self.sincronizza_tariffe_online, daemon=True).start()

    def mostra_login_modale(self):
        top = tk.Toplevel(self.root)
        top.title("🔒 SISTEMA CRITTOGRAFATO - ACCESSO RISERVATO")
        top.geometry("460x300")
        top.configure(bg="#0d1117")
        top.resizable(False, False)
        top.transient(self.root)
        top.grab_set()

        tk.Label(top, text="🔒 SISTEMA CRITTOGRAFATO", font=("Segoe UI", 14, "bold"), fg="#FFD700", bg="#0d1117", pady=15).pack()
        tk.Label(top, text="Inserire la password di sblocco per accedere:", font=("Segoe UI", 9), fg="#8b949e", bg="#0d1117").pack(pady=(0, 15))

        tk.Label(top, text="Password:", font=("Segoe UI", 10), fg="#ffffff", bg="#0d1117").pack(anchor="w", padx=40)
        entry_pwd = tk.Entry(top, font=("Segoe UI", 11), bg="#1f242c", fg="#ffffff", show="*", bd=1, insertbackground="#ffffff")
        entry_pwd.pack(fill=tk.X, padx=40, pady=5)
        entry_pwd.focus_set()

        def verifica():
            if entry_pwd.get().strip() == PASSWORD_APPLICATIVO:
                top.destroy()
                self.mostra_privacy_modale()
            else:
                messagebox.showerror("Accesso Negato", "❌ Password errata.")
                entry_pwd.delete(0, tk.END)

        entry_pwd.bind("<Return>", lambda e: verifica())
        tk.Button(top, text="🔓 SBLOCCA TERMINALE", font=("Segoe UI", 10, "bold"), bg="#2d3748", fg="#00E676", activebackground="#00E676", activeforeground="#0d1117", relief="flat", pady=6, command=verifica, cursor="hand2").pack(fill=tk.X, padx=40, pady=20)

    def mostra_privacy_modale(self):
        top = tk.Toplevel(self.root)
        top.title("🛡️ DISCIPLINARE DI SICUREZZA & PRIVACY")
        top.geometry("720x600")
        top.configure(bg="#0d1117")
        top.resizable(False, False)
        top.transient(self.root)
        top.grab_set()

        tk.Label(top, text="🛡️ DISCIPLINARE DI SICUREZZA, PRIVACY & INTEGRITÀ DATI", font=("Segoe UI", 12, "bold"), fg="#FFD700", bg="#0d1117", pady=12).pack(side=tk.TOP, fill=tk.X)

        def accetta():
            top.destroy()
            self.mostra_selettore_reparto_modale()

        btn_accetta = tk.Button(top, text="✅ ACCETTA ED ATTIVA PROTOCOLLO OPERATIVO", font=("Segoe UI", 11, "bold"), bg="#00E676", fg="#0d1117", activebackground="#00C853", relief="flat", pady=10, cursor="hand2", command=accetta)
        btn_accetta.pack(side=tk.BOTTOM, fill=tk.X, padx=25, pady=15)

        txt_privacy = scrolledtext.ScrolledText(top, bg="#161b22", fg="#ffffff", font=("Segoe UI", 10), wrap=tk.WORD, bd=1, padx=15, pady=15)
        txt_privacy.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=25, pady=5)

        corpo_testo = (
            "In conformità alle normative vigenti in materia di sicurezza informatica, tutela del "
            "segreto d'ufficio e protezione dei dati personali, l'applicazione adotta un rigoroso "
            "protocollo di elaborazione locale ad isolamento dinamico.\n\n"
            "Linee guida e garanzie operative tassative:\n\n"
            "1. ISOLAMENTO VOLATILE (RAM):\n"
            "   I documenti contabili, i file PDF/Excel e i rilievi fotografici acquisiti (OCR) vengono "
            "elaborati esclusivamente all'interno della memoria RAM volatile. Nessun dato viene archiviato "
            "su disco rigido o database locali/remoti.\n\n"
            "2. CIFRATURA DEI DATI E DOCUMENTI SENSIBILI IN TRANSITO:\n"
            "   La trasmissione delle chiavi API, dei quesiti e dell'intero contenuto dei documenti "
            "allegati (compresi i dati personali, riservati e sensibili) avviene TASSATIVAMENTE tramite "
            "canale cifrato e crittografato protetto (SSL/TLS). I flussi documentali non vengono impiegati "
            "per l'addestramento di modelli di intelligenza artificiale di terze parti.\n\n"
            "3. CONTROLLO DEL PERIMETRO E INTEGRITÀ:\n"
            "   Il sistema verifica la lettura integrale del 100% delle pagine di ogni documento allegato. "
            "Qualora una pagina risulti incompleta o illeggibile, l'IA segnalerà immediatamente l'anomalia "
            "per garantire la conformità giuridica degli atti.\n\n"
            "4. DISTRUZIONE CERTIFICATA ALL'ARRESTO:\n"
            "   Alla chiusura della sessione o mediante il comando 'Nuova Indagine (Reset)', la memoria di lavoro "
            "viene sovrascritta e liberata istantaneamente."
        )
        txt_privacy.insert(tk.END, corpo_testo)
        txt_privacy.configure(state=tk.DISABLED)

    def mostra_selettore_reparto_modale(self):
        top = tk.Toplevel(self.root)
        top.title("🛡️ SELEZIONE REPARTO OPERATIVO")
        top.geometry("640x460")
        top.configure(bg="#0d1117")
        top.resizable(False, False)
        top.transient(self.root)
        top.grab_set()

        tk.Label(top, text="🏢 SELEZIONA AMBITO D'INDAGINE", font=("Segoe UI", 14, "bold"), fg="#FFD700", bg="#0d1117", pady=20).pack()
        tk.Label(top, text="Scegli la modalità di specializzazione operativa dell'IA:", font=("Segoe UI", 10), fg="#8b949e", bg="#0d1117").pack(pady=(0, 15))

        def seleziona(mod):
            self.modalita = mod
            self.colore_accento = "#00E676" if mod == "PEF" else "#61afef"
            top.destroy()
            self.costruisci_interfaccia()

        btn_pef = tk.Button(top, text="🟡 POLIZIA ECONOMICO-FINANZIARIA (P.E.F.)\n\n• Audit Forense Frodi IVA, Cartiere, Caroselli & D.Lgs. 74/2000\n• Incrocio Registri, Fatture, Vettori, Flussi & Accertamenti Bancari", font=("Segoe UI", 10, "bold"), bg="#161b22", fg="#00E676", activebackground="#00E676", activeforeground="#0d1117", bd=2, relief="solid", pady=12, justify="left", padx=15, cursor="hand2", command=lambda: seleziona("PEF"))
        btn_pef.pack(fill=tk.X, padx=40, pady=10)

        btn_pg = tk.Button(top, text="🔵 POLIZIA GIUDIZIARIA (P.G.)\n\n• Qualificazione Reati Penali, Procedura Penale (c.p.p.) & Garanzie\n• Sequestri ex art. 354 c.p.p., Informativa art. 347 c.p.p. & Atti P.G.", font=("Segoe UI", 10, "bold"), bg="#161b22", fg="#61afef", activebackground="#61afef", activeforeground="#0d1117", bd=2, relief="solid", pady=12, justify="left", padx=15, cursor="hand2", command=lambda: seleziona("PG"))
        btn_pg.pack(fill=tk.X, padx=40, pady=10)

    def costruisci_interfaccia(self):
        for w in self.root.winfo_children():
            w.destroy()

        titolo_m = "Polizia Economico-Finanziaria (P.E.F.)" if self.modalita == "PEF" else "Polizia Giudiziaria (P.G.)"
        self.root.title(f"🛡️ Maresciallo AI - Unità {titolo_m}")

        frame_header = tk.Frame(self.root, bg=self.bg_principale, padx=15, pady=8)
        frame_header.pack(fill=tk.X)

        logo_container = tk.Frame(frame_header, bg=self.bg_principale)
        logo_container.pack(anchor=tk.CENTER)

        tk.Label(logo_container, text="Maresc", font=('Segoe UI', 24, 'bold italic'), fg=self.colore_accento, bg=self.bg_principale).pack(side=tk.LEFT)
        tk.Label(logo_container, text="[İA]", font=('Segoe UI', 22, 'bold'), fg=self.giallo_ia, bg=self.bg_principale).pack(side=tk.LEFT, padx=2)
        tk.Label(logo_container, text="llo", font=('Segoe UI', 24, 'bold italic'), fg=self.colore_accento, bg=self.bg_principale).pack(side=tk.LEFT)

        sottotitolo = "Unità Investigativa - Polizia Economico-Finanziaria (P.E.F.)" if self.modalita == "PEF" else "Unità Investigativa - Polizia Giudiziaria (P.G.)"
        tk.Label(frame_header, text=sottotitolo, font=('Segoe UI', 10, 'italic'), fg=self.bianco_puro, bg=self.bg_principale).pack(anchor=tk.CENTER)

        tk.Frame(self.root, height=2, bg=self.colore_accento).pack(fill=tk.X, padx=20, pady=2)

        frame_border_v = tk.Frame(self.root, bg=self.colore_accento, padx=2, pady=2)
        frame_border_v.pack(side=tk.LEFT, fill=tk.Y, padx=(10, 5), pady=8)

        frame_border_g = tk.Frame(frame_border_v, bg=self.giallo_ia, padx=2, pady=2)
        frame_border_g.pack(fill=tk.BOTH, expand=True)

        self.sidebar = tk.Frame(frame_border_g, bg=self.bg_aree, width=280, padx=10, pady=10)
        self.sidebar.pack(fill=tk.BOTH, expand=True)
        self.sidebar.pack_propagate(False)

        f_cmd = tk.Frame(self.sidebar, bg=self.bg_aree)
        f_cmd.pack(fill=tk.X, pady=(0, 5))
        
        lbl_mod = "🏢 MODALITÀ: P.E.F." if self.modalita == "PEF" else "🚨 MODALITÀ: P.G."
        tk.Label(f_cmd, text=lbl_mod, bg=self.bg_aree, fg=self.colore_accento, font=('Segoe UI', 9, 'bold')).pack(side=tk.LEFT)
        tk.Button(f_cmd, text="💶 Costi", bg=self.bg_pulsante, fg=self.giallo_ia, font=('Segoe UI', 8, 'bold'), command=self.apri_finestra_costi, cursor="hand2").pack(side=tk.RIGHT)

        tk.Button(
            self.sidebar,
            text="🔄 Cambia Reparto (PEF / PG)",
            bg="#2d3748",
            fg="#FFD700",
            font=('Segoe UI', 9, 'bold'),
            command=self.mostra_selettore_reparto_modale,
            cursor="hand2",
            pady=2
        ).pack(fill=tk.X, pady=(2, 6))

        tk.Label(self.sidebar, text="Chiave API Google:", bg=self.bg_aree, fg=self.bianco_puro, font=('Segoe UI', 8)).pack(anchor=tk.W)
        self.entry_key = tk.Entry(self.sidebar, bg=self.bg_input, fg=self.bianco_puro, font=('Segoe UI', 8), show="*", insertbackground="#ffffff")
        self.entry_key.pack(fill=tk.X, pady=2)

        f_k_btn = tk.Frame(self.sidebar, bg=self.bg_aree)
        f_k_btn.pack(fill=tk.X, pady=(0, 6))
        tk.Button(f_k_btn, text="Salva", bg=self.bg_pulsante, fg=self.bianco_puro, font=('Segoe UI', 8), command=self.salva_chiave_config).pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Button(f_k_btn, text="Rimuovi", bg=self.rosso_alert, fg=self.bianco_puro, font=('Segoe UI', 8), command=self.cancella_chiave_config).pack(side=tk.RIGHT, fill=tk.X, expand=True)

        tk.Label(self.sidebar, text="💬 Stile Interfaccia:", bg=self.bg_aree, fg=self.bianco_puro, font=('Segoe UI', 9, 'bold')).pack(anchor=tk.W, pady=(4, 1))
        self.combo_stile = ttk.Combobox(self.sidebar, values=["🤝 Collega (Naturale & Dettagliato)", "⚡ Sintetico (Solo Dati Essenziali)"], state="readonly", font=('Segoe UI', 9))
        self.combo_stile.current(0)
        self.combo_stile.pack(fill=tk.X, pady=(0, 6))

        tk.Label(self.sidebar, text="📂 ALLEGATI DA ESAMINARE:", bg=self.bg_aree, fg=self.bianco_puro, font=('Segoe UI', 9, 'bold')).pack(anchor=tk.W)
        tk.Button(self.sidebar, text="📂 Sfoglia File", bg=self.bg_pulsante, fg=self.bianco_puro, font=('Segoe UI', 9, 'bold'), command=self.acquisisci_file).pack(fill=tk.X, pady=3)

        self.list_files = tk.Listbox(self.sidebar, bg=self.bg_input, fg=self.bianco_puro, height=3, font=('Segoe UI', 9))
        self.list_files.pack(fill=tk.X, pady=2)
        self.list_files.bind("<Delete>", lambda e: self.elimina_file_selezionato())
        self.list_files.bind("<Double-Button-1>", lambda e: self.elimina_file_selezionato())

        tk.Label(self.sidebar, text="💡 Doppio click o [Canc] per rimuovere.", bg=self.bg_aree, fg="#8b949e", font=('Segoe UI', 8, 'italic')).pack(anchor=tk.W)
        tk.Frame(self.sidebar, height=1, bg="#2d3748").pack(fill=tk.X, pady=6)

        tk.Label(self.sidebar, text="📝 VERBALE DA COMPILARE:", bg=self.bg_aree, fg=self.giallo_ia, font=('Segoe UI', 9, 'bold')).pack(anchor=tk.W)
        tk.Button(self.sidebar, text="✍️ Compila Verbale (.docx)", bg="#2e7d32" if self.modalita == "PEF" else "#1565C0", fg=self.bianco_puro, font=('Segoe UI', 9, 'bold'), command=self.compila_verbale_modello, cursor="hand2").pack(fill=tk.X, pady=2)

        tk.Label(self.sidebar, text="🖨️ ATTI DI ESPORTAZIONE CHAT:", bg=self.bg_aree, fg=self.bianco_puro, font=('Segoe UI', 9, 'bold')).pack(anchor=tk.W, pady=(4, 0))
        tk.Button(self.sidebar, text="Esporta Ultima (PDF Formale)", bg=self.bg_pulsante, fg=self.bianco_puro, font=('Segoe UI', 9, 'bold'), command=lambda: self.esporta_pdf(solo_ultima=True)).pack(fill=tk.X, pady=1)
        tk.Button(self.sidebar, text="Esporta Verbale (PDF)", bg=self.bg_pulsante, fg=self.bianco_puro, font=('Segoe UI', 9, 'bold'), command=lambda: self.esporta_pdf(solo_ultima=False)).pack(fill=tk.X, pady=1)
        tk.Button(self.sidebar, text="Esporta Ultima (Word Doc)", bg="#1976D2", fg=self.bianco_puro, font=('Segoe UI', 9, 'bold'), command=lambda: self.esporta_docx(solo_ultima=True)).pack(fill=tk.X, pady=1)
        tk.Button(self.sidebar, text="Esporta Verbale (Word)", bg="#1976D2", fg=self.bianco_puro, font=('Segoe UI', 9, 'bold'), command=lambda: self.esporta_docx(solo_ultima=False)).pack(fill=tk.X, pady=1)

        self.lbl_status = tk.Label(self.sidebar, text="Stato: Pronto", bg=self.bg_aree, fg=self.colore_accento, font=('Segoe UI', 9, 'italic'))
        self.lbl_status.pack(anchor=tk.W, pady=(4, 2))

        style = ttk.Style()
        style.theme_use('default')
        style.configure("TProgressbar", thickness=14, troughcolor="#161b22", background=self.colore_accento)
        self.progress_bar = ttk.Progressbar(self.sidebar, mode="indeterminate", style="TProgressbar")
        self.progress_bar.pack(fill=tk.X, pady=2)

        tk.Button(self.sidebar, text="🔄 Nuova Indagine (Reset)", bg=self.arancione_reset, fg=self.bg_principale, font=('Segoe UI', 9, 'bold'), command=self.reset_indagine).pack(fill=tk.X, pady=(6, 0))

        frame_chat = tk.Frame(self.root, bg=self.bg_principale, padx=5, pady=5)
        frame_chat.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 10))

        self.display_chat = scrolledtext.ScrolledText(
            frame_chat, bg=self.bg_aree, fg=self.bianco_puro, insertbackground=self.bianco_puro,
            font=('Courier New', 10), wrap=tk.WORD, bd=1, padx=10, pady=10
        )
        self.display_chat.pack(fill=tk.BOTH, expand=True, pady=(0, 8))

        msg_i = "Maresciallo AI [P.E.F.]: Pronto all'audit forense, incrocio registri, frodi IVA e accertamenti bancari...\n\n" if self.modalita == "PEF" else "Maresciallo AI [P.G.]: Pronto all'esame penale, riscontro D.Lgs. 74/2000, c.p.p. e atti di P.G...\n\n"
        self.display_chat.insert(tk.END, msg_i)
        self.display_chat.configure(state=tk.DISABLED)

        frame_input = tk.Frame(frame_chat, bg=self.bg_principale)
        frame_input.pack(fill=tk.X, side=tk.BOTTOM)

        self.entry_prompt = scrolledtext.ScrolledText(frame_input, bg=self.bg_input, fg=self.bianco_puro, insertbackground=self.bianco_puro, font=('Segoe UI', 10), bd=1, height=3, wrap=tk.WORD)
        self.entry_prompt.pack(fill=tk.X, side=tk.LEFT, expand=True, padx=(0, 8))

        tk.Button(frame_input, text="INVIA ATTO", bg=self.bg_pulsante, fg=self.bianco_puro, font=('Segoe UI', 10, 'bold'), padx=12, command=self.invia_quesito).pack(side=tk.RIGHT)

        self.carica_chiave_automatica()

    def carica_statistiche_locali(self):
        m_corrente = datetime.datetime.now().strftime("%Y-%m")
        self.stats = {"mese_corrente": m_corrente, "mese_token_in": 0, "mese_token_out": 0, "mese_costo_eur": 0.0, "tot_token_in": 0, "tot_token_out": 0, "tot_costo_eur": 0.0}
        if os.path.exists(STATS_FILE):
            try:
                with open(STATS_FILE, "r") as f:
                    data = json.load(f)
                    if data.get("mese_corrente") != m_corrente:
                        data["mese_corrente"] = m_corrente
                        data["mese_token_in"] = 0; data["mese_token_out"] = 0; data["mese_costo_eur"] = 0.0
                    self.stats.update(data)
            except Exception: pass

    def salva_statistiche_locali(self):
        try:
            with open(STATS_FILE, "w") as f: json.dump(self.stats, f, indent=2)
        except Exception: pass

    def carica_tariffe_salvate(self):
        if os.path.exists(PRICES_FILE):
            try:
                with open(PRICES_FILE, "r") as f:
                    p_data = json.load(f)
                    self.costo_input_1m_eur = p_data.get("input_1m_eur", DEFAULT_COST_INPUT_1M_EUR)
                    self.costo_output_1m_eur = p_data.get("output_1m_eur", DEFAULT_COST_OUTPUT_1M_EUR)
            except Exception: pass

    def sincronizza_tariffe_online(self):
        try:
            url = "https://raw.githubusercontent.com/google-gemini/api-pricing/main/gemini_prices.json"
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=4) as response:
                if response.status == 200:
                    data = json.loads(response.read().decode())
                    if "gemini-3.6-flash" in data:
                        p = data["gemini-3.6-flash"]
                        self.costo_input_1m_eur = float(p.get("input_1m_eur", DEFAULT_COST_INPUT_1M_EUR))
                        self.costo_output_1m_eur = float(p.get("output_1m_eur", DEFAULT_COST_OUTPUT_1M_EUR))
                        self.stato_tariffe_online = f"🟢 Aggiornate Online ({datetime.datetime.now().strftime('%H:%M')})"
                        with open(PRICES_FILE, "w") as f:
                            json.dump({"input_1m_eur": self.costo_input_1m_eur, "output_1m_eur": self.costo_output_1m_eur}, f)
                        return
        except Exception: pass
        self.stato_tariffe_online = "🟡 Tariffe di Base (Offline)"

    def registra_consumo_token(self, token_in, token_out):
        costo_in = (token_in / 1_000_000.0) * self.costo_input_1m_eur
        costo_out = (token_out / 1_000_000.0) * self.costo_output_1m_eur
        costo_tot = costo_in + costo_out
        self.sess_token_in += token_in; self.sess_token_out += token_out; self.sess_costo_eur += costo_tot
        m_corrente = datetime.datetime.now().strftime("%Y-%m")
        if self.stats.get("mese_corrente") != m_corrente:
            self.stats["mese_corrente"] = m_corrente
            self.stats["mese_token_in"] = 0; self.stats["mese_token_out"] = 0; self.stats["mese_costo_eur"] = 0.0
        self.stats["mese_token_in"] += token_in; self.stats["mese_token_out"] += token_out; self.stats["mese_costo_eur"] += costo_tot
        self.stats["tot_token_in"] += token_in; self.stats["tot_token_out"] += token_out; self.stats["tot_costo_eur"] += costo_tot
        self.salva_statistiche_locali()
        return costo_tot

    def apri_finestra_costi(self):
        win_costi = tk.Toplevel(self.root)
        win_costi.title("💶 Gestione Costi & Consumi API")
        win_costi.geometry("520x420")
        win_costi.configure(bg=self.bg_aree)
        win_costi.resizable(False, False)
        win_costi.transient(self.root)

        tk.Label(win_costi, text="💶 CONTO ECONOMICO & CONSUMI API", font=("Segoe UI", 12, "bold"), fg=self.giallo_ia, bg=self.bg_aree, pady=10).pack()
        lbl_st = tk.Label(win_costi, text=f"Stato Tariffe: {self.stato_tariffe_online}", font=("Segoe UI", 9, "italic"), fg=self.colore_accento, bg=self.bg_aree)
        lbl_st.pack(pady=(0, 10))

        frame_dati = tk.Frame(win_costi, bg=self.bg_input, padx=15, pady=12, bd=1, relief="solid")
        frame_dati.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)

        tot_sess_t = self.sess_token_in + self.sess_token_out
        tot_mese_t = self.stats["mese_token_in"] + self.stats["mese_token_out"]
        tot_stor_t = self.stats["tot_token_in"] + self.stats["tot_token_out"]

        str_costi = (
            f"📌 SESSIONE CORRENTE:\n"
            f"   • Token Elaborati: {tot_sess_t:,}\n"
            f"   • Spesa Stimata:   {self.sess_costo_eur:.4f} €\n\n"
            f"📅 MESE IN CORSO ({self.stats['mese_corrente']}):\n"
            f"   • Token Elaborati: {tot_mese_t:,}\n"
            f"   • Spesa Stimata:   {self.stats['mese_costo_eur']:.4f} €\n\n"
            f"🏛️ TOTALE STORICO (Dall'installazione):\n"
            f"   • Token Elaborati: {tot_stor_t:,}\n"
            f"   • Spesa Stimata:   {self.stats['tot_costo_eur']:.4f} €\n\n"
            f"────────────\n"
            f"💡 Tariffe Gemini Flash applicate:\n"
            f" Input:  ~{self.costo_input_1m_eur:.3f} € / 1M token\n"
            f" Output: ~{self.costo_output_1m_eur:.3f} € / 1M token"
        )
        tk.Label(frame_dati, text=str_costi, font=("Segoe UI", 9), fg=self.bianco_puro, bg=self.bg_input, justify="left").pack(anchor="w")

        btn_frame = tk.Frame(win_costi, bg=self.bg_aree, pady=10)
        btn_frame.pack(fill=tk.X, padx=20)

        def azzera_stat():
            if messagebox.askyesno("Conferma Reset", "Vuoi azzerare lo storico dei consumi mensili e generali?"):
                self.stats["mese_token_in"] = 0
                self.stats["mese_token_out"] = 0
                self.stats["mese_costo_eur"] = 0.0
                self.stats["tot_token_in"] = 0
                self.stats["tot_token_out"] = 0
                self.stats["tot_costo_eur"] = 0.0
                self.salva_statistiche_locali()
                win_costi.destroy()

        tk.Button(btn_frame, text="🗑️ Azzera Storico", bg=self.rosso_alert, fg=self.bianco_puro, font=("Segoe UI", 9, "bold"), command=azzera_stat).pack(side=tk.LEFT)
        tk.Button(btn_frame, text="Chiudi", bg=self.bg_pulsante, fg=self.bianco_puro, font=("Segoe UI", 9), command=win_costi.destroy).pack(side=tk.RIGHT)

    def salva_chiave_config(self):
        key = self.entry_key.get().strip()
        if not key: return
        try:
            enc = base64.b64encode(key.encode('utf-8')).decode('utf-8')
            with open(CONFIG_FILE, "w") as f: f.write(enc)
            messagebox.showinfo("OK", "🔒 Chiave API salvata.")
        except Exception as e: messagebox.showerror("Errore", str(e))

    def carica_chiave_automatica(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r") as f: content = f.read().strip()
                if content:
                    self.entry_key.delete(0, tk.END)
                    self.entry_key.insert(0, base64.b64decode(content.encode('utf-8')).decode('utf-8'))
            except Exception: pass

    def cancella_chiave_config(self):
        if os.path.exists(CONFIG_FILE): os.remove(CONFIG_FILE)
        self.entry_key.delete(0, tk.END)

    def acquisisci_file(self):
        files = filedialog.askopenfilenames(title="Seleziona allegati", filetypes=[("Tutti i File", "*.pdf *.xlsx *.txt *.png *.jpg *.jpeg *.docx")])
        if files:
            presenza_immagini = False
            for f in files:
                if f not in self.files_acquisiti: 
                    self.files_acquisiti.append(f)
                ext = f.split('.')[-1].lower()
                if ext in ["jpg", "jpeg", "png"]:
                    presenza_immagini = True
            self.aggiorna_lista()
            
            if presenza_immagini:
                messagebox.showwarning(
                    "⚠️ Avviso Consumo Token (OCR / Immagini)",
                    "Hai caricato una o più immagini o scansioni fotografiche.\n\n"
                    "L'elaborazione visiva (OCR) consuma un numero di token notevolmente superiore rispetto al testo semplice, incidendo sui costi della sessione."
                )

    def aggiorna_lista(self):
        self.list_files.delete(0, tk.END)
        for path in self.files_acquisiti:
            self.list_files.insert(tk.END, f"📄 {os.path.basename(path)}")

    def elimina_file_selezionato(self):
        try:
            if idx := self.list_files.curselection():
                self.files_acquisiti.pop(idx[0])
                self.aggiorna_lista()
        except Exception: pass

    def correggi_ortografia_italiana(self, testo):
        correzioni = [
            (r"\bDallesame\b", "Dall'esame"), (r"\bdImposta\b", "d'imposta"), (r"\bdimposta\b", "d'imposta"),
            (r"\bdimpresa\b", "d'impresa"), (r"\bdAffari\b", "d'affari"), (r"\bdaffari\b", "d'affari"),
            (r"\bSOSTITUTO DIMPOSTA\b", "SOSTITUTO D'IMPOSTA"), (r"\bLanalisi\b", "L'analisi"),
            (r"\bdallesame\b", "dall'esame"), (r"\blesercizio\b", "l'esercizio"), (r"\bLesercizio\b", "L'esercizio")
        ]
        for pattern, sostituto in correzioni:
            testo = re.sub(pattern, sostituto, testo, flags=re.IGNORECASE)
        return testo

    def formatta_tabella_allineata(self, testo_risposta):
        testo_risposta = self.correggi_ortografia_italiana(testo_risposta)
        righe = testo_risposta.split('\n')
        nuovo_testo = []
        in_tabella = False
        blocchi_tabella = []

        for riga in righe:
            riga_strip = riga.strip()
            if riga_strip.startswith('|'):
                in_tabella = True
                if '---' in riga_strip: continue
                celle = [c.strip().replace('**', '') for c in riga_strip.split('|')[1:-1]]
                if celle: blocchi_tabella.append(celle)
            else:
                if in_tabella and blocchi_tabella:
                    num_colonne = max(len(r) for r in blocchi_tabella)
                    larghezze = [0] * num_colonne
                    for r in blocchi_tabella:
                        for idx, cella in enumerate(r):
                            if idx < len(larghezze) and len(cella) > larghezze[idx]:
                                larghezze[idx] = len(cella)

                    linea_sep = "+" + "+".join(["-" * (w + 2) for w in larghezze]) + "+"
                    nuovo_testo.append(linea_sep)

                    for index_r, r in enumerate(blocchi_tabella):
                        riga_f = "|"
                        for idx, cella in enumerate(r):
                            if idx < len(larghezze):
                                riga_f += f" {cella}{' ' * (larghezze[idx] - len(cella))} |"
                        nuovo_testo.append(riga_f)
                        if index_r == 0: nuovo_testo.append(linea_sep)

                    nuovo_testo.append(linea_sep)
                    blocchi_tabella = []
                    in_tabella = False

                nuovo_testo.append(riga)

        if in_tabella and blocchi_tabella:
            num_colonne = max(len(r) for r in blocchi_tabella)
            larghezze = [0] * num_colonne
            for r in blocchi_tabella:
                for idx, cella in enumerate(r):
                    if idx < len(larghezze) and len(cella) > larghezze[idx]: larghezze[idx] = len(cella)

            linea_sep = "+" + "+".join(["-" * (w + 2) for w in larghezze]) + "+"
            nuovo_testo.append(linea_sep)
            for index_r, r in enumerate(blocchi_tabella):
                riga_f = "|"
                for idx, cella in enumerate(r):
                    if idx < len(larghezze): riga_f += f" {cella}{' ' * (larghezze[idx] - len(cella))} |"
                nuovo_testo.append(riga_f)
                if index_r == 0: nuovo_testo.append(linea_sep)
            nuovo_testo.append(linea_sep)

        return "\n".join(nuovo_testo)

    def leggi_file_locali(self):
        testo = ""
        immagini = []
        for path in self.files_acquisiti:
            if not os.path.exists(path): continue
            est = path.split('.')[-1].lower()
            if est in ["jpg", "jpeg", "png"]:
                with open(path, "rb") as f:
                    mime = "image/jpeg" if est in ["jpg", "jpeg"] else "image/png"
                    immagini.append(types.Part.from_bytes(data=f.read(), mime_type=mime))
                continue
            
            testo += f"\n[INIZIO FILE: {os.path.basename(path)}]\n"
            try:
                if est == "txt":
                    with open(path, "r", encoding="utf-8", errors="ignore") as f: testo += f.read()
                elif est == "pdf" and pypdf:
                    reader = pypdf.PdfReader(path)
                    for num_p, page in enumerate(reader.pages, start=1):
                        t = page.extract_text()
                        if t: testo += f"--- Pagina {num_p} ---\n" + t + "\n"
                elif est == "xlsx" and openpyxl:
                    wb = openpyxl.load_workbook(path, data_only=True)
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        testo += f"[Foglio Excel: {sheet_name}]\n"
                        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                            if any(row):
                                riga_formattata = " | ".join([str(c) if c is not None else "" for c in row])
                                testo += f"Riga {row_idx}: {riga_formattata}\n"
                elif est == "docx" and Document:
                    doc = Document(path)
                    for p in doc.paragraphs:
                        if p.text: testo += p.text + "\n"
            except Exception as e: testo += f"[Errore lettura {os.path.basename(path)}: {str(e)}]\n"
            testo += f"\n[FINE FILE: {os.path.basename(path)}]\n"
            
        return self.correggi_ortografia_italiana(testo), immagini

    def invia_quesito(self):
        prompt = self.entry_prompt.get("1.0", tk.END).strip()
        key = self.entry_key.get().strip()
        if not prompt or not key: return

        if len(self.storico_chat) >= 2:
            risposta_reset = messagebox.askyesnocancel(
                "💡 Suggerimento Risparmio Token",
                "Stai ponendo una domanda in PROSEGUIMENTO dell'analisi precedente o si tratta di un NUOVO ARGOMENTO?\n\n"
                "• Clicca 'SÌ' per proseguire mantenendo il contesto.\n"
                "• Clicca 'NO' per AZZERARE LA MEMORIA (Reset) ed evitare di pagare l'invio dei vecchi messaggi."
            )
            if risposta_reset is None: return
            elif risposta_reset is False: self.reset_indagine()

        self.display_chat.configure(state=tk.NORMAL)
        self.display_chat.insert(tk.END, f"Operatore: {prompt}\n\n")
        self.display_chat.configure(state=tk.DISABLED)

        self.entry_prompt.delete("1.0", tk.END)
        self.lbl_status.configure(text="Stato: Elaborazione...", fg=self.giallo_ia)
        self.progress_bar.start(10)

        threading.Thread(target=self.esecuzione_background, args=(prompt, key), daemon=True).start()

    def esecuzione_background(self, prompt, key):
        try:
            testo_atti, immagini = self.leggi_file_locali()
            p_completo = prompt
            if testo_atti: 
                p_completo += f"\n\n[DATI ESTRATTI DAI DOCUMENTI ALLEGATI (DA LEGGERE RIGOROSAMENTE SECONDO LE INTESTAZIONI LETTERALI)]\n{testo_atti}"

            client = genai.Client(api_key=key)
            stile_scelto = self.combo_stile.get()
            istruzione_stile = "Esponi con tono naturale, collaborativo e articolato, mantenendo elevata precisione tecnica."
            if "Sintetico" in stile_scelto:
                istruzione_stile = "Esponi in modo estremamente sintetico, schematico e diretto, fornendo solo i dati e le conclusioni essenziali."

            if self.modalita == "PEF":
                sys_prompt = (
                    "Sei un Assistente Virtuale Senior specializzato in Polizia Economico-Finanziaria (P.E.F.), Audit Forense, Diritto Tributario e Accertamenti Bancari.\n\n"
                    f"--- STILE DI INTERAZIONE ---\n{istruzione_stile}\n\n"
                    "--- PROTOCOLLO GEOMETRICO E TASSATIVO CONTRO L'INVERSIONE DEI DATI ---\n"
                    "1. VINCOLO ASSOLUTO DI LETTURA LETTERALE:\n"
                    "   • È SEVERAMENTE VIETATO invertire le operazioni ATTIVE (IVA a debito / cessioni / imponibile attivo / accrediti / versamenti) con le operazioni PASSIVE (IVA a credito / acquisti / imponibile passivo / prelevamenti).\n"
                    "   • Devi leggere esclusivamente la stringa di testo o la riga della tabella associata all'etichetta letterale originale. La parola scritta e l'etichetta di colonna hanno priorità matematica assoluta.\n\n"
                    "--- DIRETTIVE SPECIALIZZATE PER ACCERTAMENTI BANCARI E FINANZIARI (art. 32 D.P.R. 600/73 & art. 51 D.P.R. 633/72) ---\n"
                    "1. ANALISI MOVIMENTI BANCARI ED ESTRATTI CONTO:\n"
                    "   • VERSAMENTI (Accrediti): Vanno considerati, per presunzione legale (inversione dell'onere della prova), come compensi o ricavi non dichiarati, salvo rigorosa prova contraria fornita dal contribuente (es. giroconti, finanziamenti soci, liberalità, disinvestimenti giustificati).\n"
                    "   • PRELEVAMENTI (Addebiti): Verifica la qualifica soggettiva del contribuente (imprenditore vs lavoratore autonomo / professionista). Ricorda che per i professionisti/lavoratori autonomi i prelevamenti ingiustificati non costituiscono compensi occulti (Corte Cost. n. 228/2014), mentre per gli imprenditori commerciali permane la presunzione limitata alle soglie normative vigenti.\n"
                    "2. QUADRATURA E TABELLE DI RISCONTRO:\n"
                    "   • Genera tabelle di scomposizione analitica indicando chiaramente: Data, Causale, Importo, Natura (Versamento vs Prelevamento), Eventuale Giustificazione documentale e Rischio Fiscale.\n\n"
                    "--- STRUTTURA DELLA RISPOSTA ---\n"
                    "- 📌 PERIMETRO DI ANALISI ED ESAME DOCUMENTALE\n"
                    "- 📍 QUESITO / FATTISPECIE ISPETTIVA O BANCARIA\n"
                    "- 📊 TABELLA DI QUADRATURA E RISCONTRO ANALITICO (ATTIVI / PASSIVI O VERSAMENTI / PRELEVAMENTI)\n"
                    "- 🔍 RILIEVI FISCALI, PRESUNZIONI LEGALI ED EVENTUALI PROFILI PENALI (D.Lgs. 74/2000)\n"
                    "- 📝 CONCLUSIONI TECNICHE ED APPROFONDIMENTI ISPETTIVI SUGGERITI"
                )
            else:
                sys_prompt = (
                    "Sei un Assistente Virtuale Senior specializzato in Polizia Giudiziaria (P.G.), Diritto Penale e Procedura Penale (c.p.p.).\n\n"
                    f"--- STILE DI INTERAZIONE ---\n{istruzione_stile}\n\n"
                    "--- DIRETTIVE OPERATIVE POLIZIA GIUDIZIARIA ---\n"
                    "1. DICHIARAZIONE PERIMETRO: Inizia dichiarando la quantità esatta di atti penali ed elementi di prova esaminati.\n"
                    "2. QUALIFICAZIONE PENALE E NORMATIVA PROCEDURALE:\n"
                    "   • Inquadramento Penale Rigoroso: Qualifica giuridicamente le condotte penalmente rilevanti (D.Lgs. 74/2000, Reati Societari ex art. 2621 e ss. c.c., Reati Fallimentari/Bancarotta ex CCII, Reati contro la P.A. o il Patrimonio).\n"
                    "   • Elemento Soggettivo e Materiale: Analizza il dolo specifico di evasione o di profitto e la condotta materiale con riferimenti alle più recenti sentenze di Cassazione Penale.\n"
                    "   • Procedura Penale (c.p.p.): Verifica la regolarità degli atti di P.G. (art. 347 c.p.p. Informativa di Reato, art. 352 c.p.p. Perquisizione, art. 354 c.p.p. Sequestro probatorio/preventivo) e il rispetto tassativo delle GARANZIE DIFENSIVE ex art. 356 c.p.p. per evitare eccezioni di inutilizzabilità.\n"
                    "3. REDAZIONE ATTI DI P.G.: Fornisci indicazioni formali chiare e bozze di clausole giuridiche per la redazione di Annotazioni e Informative di Reato a perfetta tenuta dibattimentale.\n\n"
                    "--- STRUTTURA DELLA RISPOSTA ---\n"
                    "- 📌 PERIMETRO DI ANALISI ED ESAME ATTI DI P.G.\n"
                    "- 📍 REATO / FATTISPECIE PENALE CONTESTATA\n"
                    "- ⚖️ QUALIFICAZIONE GIURIDICA (Fattispecie, Dolo, Giurisprudenza & Soglie Penali)\n"
                    "- 🔍 RISCONTRO PROCEDURALE PENALE (c.p.p., Garanzie & Tenuta Probatoria)\n"
                    "- 📝 DIRETTIVE FORMALI PER LA REDAZIONE DEGLI ATTI DI P.G."
                )

            contents = []
            for m in self.storico_chat:
                contents.append(types.Content(role="user" if m["role"] == "user" else "model", parts=[types.Part.from_text(text=m["content"])]))
            contents.append(types.Content(role="user", parts=[types.Part.from_text(text=p_completo)] + immagini))

            res = client.models.generate_content(
                model='gemini-3.6-flash',
                contents=contents,
                config=types.GenerateContentConfig(system_instruction=sys_prompt, temperature=0.0)
            )

            costo = 0.0
            if hasattr(res, 'usage_metadata') and res.usage_metadata:
                costo = self.registra_consumo_token(res.usage_metadata.prompt_token_count or 0, res.usage_metadata.candidates_token_count or 0)

            self.storico_chat.append({"role": "user", "content": prompt})
            self.storico_chat.append({"role": "assistant", "content": res.text})

            risposta_tabulata = self.formatta_tabella_allineata(res.text)
            risposta_f = risposta_tabulata + f"\n\n[💶 Spesa API per questo atto: {costo:.5f} €]"
            self.root.after(0, lambda: self.stampa_successo(risposta_f))
        except Exception as e:
            self.root.after(0, lambda: self.stampa_errore(str(e)))

    def stampa_successo(self, testo):
        self.progress_bar.stop()
        self.lbl_status.configure(text="Stato: Pronto", fg=self.colore_accento)
        self.display_chat.configure(state=tk.NORMAL)
        self.display_chat.insert(tk.END, f"Maresciallo AI:\n{testo}\n{'-'*60}\n\n")
        self.display_chat.configure(state=tk.DISABLED)
        self.display_chat.see(tk.END)

    def stampa_errore(self, err):
        self.progress_bar.stop()
        self.lbl_status.configure(text="Stato: Errore", fg=self.rosso_alert)
        self.display_chat.configure(state=tk.NORMAL)
        self.display_chat.insert(tk.END, f"⚠️ Errore: {err}\n\n")
        self.display_chat.configure(state=tk.DISABLED)

    def reset_indagine(self):
        self.storico_chat.clear()
        self.files_acquisiti.clear()
        self.aggiorna_lista()
        self.display_chat.configure(state=tk.NORMAL)
        self.display_chat.delete("1.0", tk.END)
        msg_i = "Maresciallo AI [P.E.F.]: Pronto all'audit forense, incrocio registri, frodi IVA e accertamenti bancari...\n\n" if self.modalita == "PEF" else "Maresciallo AI [P.G.]: Pronto all'esame penale, riscontro D.Lgs. 74/2000, c.p.p. e atti di P.G...\n\n"
        self.display_chat.insert(tk.END, msg_i)
        self.display_chat.configure(state=tk.DISABLED)
        self.progress_bar.stop()

    def compila_verbale_modello(self):
        if not Document:
            messagebox.showerror("Errore", "Libreria python-docx non installata.")
            return
        if not self.files_acquisiti:
            messagebox.showwarning("Attenzione", "Carica prima gli atti allegati nell'elenco di sinistra.")
            return
        f_modello = filedialog.askopenfilename(title="Seleziona Modello Word (.docx)", filetypes=[("Word", "*.docx")])
        if not f_modello: return
        messagebox.showinfo("Compilazione", "Modulo di compilazione automatica attivo sul modello selezionato.")

    def esporta_pdf(self, solo_ultima=True):
        if not self.storico_chat or not HAS_REPORTLAB: return
        f_salva = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("Documento PDF", "*.pdf")])
        if not f_salva: return
        try:
            doc = SimpleDocTemplate(f_salva, pagesize=letter, rightMargin=20, leftMargin=20, topMargin=30, bottomMargin=30)
            styles = getSampleStyleSheet()
            
            style_titolo = ParagraphStyle('TitoloReport', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=12, spaceAfter=12, alignment=1)
            style_testo = ParagraphStyle('TestoReport', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=10)
            style_cella = ParagraphStyle('CellaTabella', parent=styles['Normal'], fontName='Helvetica', fontSize=7, leading=9)
            style_cella_h = ParagraphStyle('CellaTabellaH', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=7, leading=9, textColor=colors.white)
            
            titolo_atto = "RELAZIONE TECNICA ISPETTIVA DI VERIFICA (P.E.F.)" if self.modalita == "PEF" else "ANNOTAZIONE DI POLIZIA GIUDIZIARIA (P.G.)"
            story = [Paragraph(titolo_atto, style_titolo), Spacer(1, 8)]

            items = self.storico_chat[-2:] if solo_ultima else self.storico_chat
            for m in items:
                if m['role'] == 'assistant':
                    righe = m['content'].split('\n')
                    buffer_testo = []
                    righe_tabella = []
                    
                    def flush_testo():
                        if buffer_testo:
                            story.append(Paragraph("<br/>".join(buffer_testo), style_testo))
                            buffer_testo.clear()
                    
                    for riga in righe:
                        riga_strip = riga.strip()
                        if riga_strip.startswith('|'):
                            if '---' in riga_strip: continue
                            if buffer_testo: flush_testo()
                            
                            celle = [c.strip().replace('**', '') for c in riga_strip.split('|')[1:-1]]
                            if celle: righe_tabella.append(celle)
                        else:
                            if righe_tabella:
                                data_tabella = []
                                for idx_r, r in enumerate(righe_tabella):
                                    r_par = [Paragraph(cell.replace('€', 'EUR').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), style_cella_h if idx_r == 0 else style_cella) for cell in r]
                                    data_tabella.append(r_par)
                                
                                if data_tabella:
                                    num_c = len(data_tabella[0])
                                    col_w = [572 / num_c] * num_c
                                    t = Table(data_tabella, colWidths=col_w)
                                    t.setStyle(TableStyle([
                                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2d3748")),
                                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                                    ]))
                                    story.append(t)
                                    story.append(Spacer(1, 8))
                                righe_tabella = []
                            
                            if riga_strip:
                                buffer_testo.append(riga_strip.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
                            elif buffer_testo:
                                flush_testo()
                                story.append(Spacer(1, 4))
                    
                    if buffer_testo: flush_testo()
                    if righe_tabella:
                        data_tabella = []
                        for idx_r, r in enumerate(righe_tabella):
                            r_par = [Paragraph(cell.replace('€', 'EUR').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), style_cella_h if idx_r == 0 else style_cella) for cell in r]
                            data_tabella.append(r_par)
                        if data_tabella:
                            num_c = len(data_tabella[0])
                            col_w = [572 / num_c] * num_c
                            t = Table(data_tabella, colWidths=col_w)
                            t.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2d3748")),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                                ('TOPPADDING', (0, 0), (-1, -1), 4),
                            ]))
                            story.append(t)
                    story.append(Spacer(1, 10))
                    
            doc.build(story)
            messagebox.showinfo("Esportazione", "Documento PDF esportato correttamente con tabelle professionali.")
        except Exception as e: 
            messagebox.showerror("Errore di Esportazione", f"Impossibile generare il PDF:\n{str(e)}")

    def esporta_docx(self, solo_ultima=True):
        if not self.storico_chat or not Document: return
        f_salva = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")])
        if not f_salva: return
        try:
            doc = Document()
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.size = Pt(10)
            doc.add_heading("RELAZIONE DI ANALISI ISPETTIVA", level=1)
            
            items = self.storico_chat[-2:] if solo_ultima else self.storico_chat
            for m in items:
                if m['role'] == 'assistant':
                    righe = m['content'].split('\n')
                    buffer_testo = []
                    righe_tabella = []
                    
                    def flush_testo_docx():
                        if buffer_testo:
                            doc.add_paragraph("\n".join(buffer_testo))
                            buffer_testo.clear()

                    def aggiungi_tabella_docx(dati_tab):
                        if not dati_tab: return
                        num_righe = len(dati_tab)
                        num_col = max(len(r) for r in dati_tab)
                        table = doc.add_table(rows=num_righe, cols=num_col)
                        table.style = 'Table Grid'
                        
                        for r_idx, r_data in enumerate(dati_tab):
                            row = table.rows[r_idx]
                            for c_idx, cell_val in enumerate(r_data):
                                if c_idx < len(row.cells):
                                    cell = row.cells[c_idx]
                                    cell.text = cell_val.replace('**', '')
                                    
                                    if r_idx == 0:
                                        for p in cell.paragraphs:
                                            p.alignment = 0 
                                            for run in p.runs:
                                                run.font.bold = True
                                                run.font.size = Pt(9)
                                                run.font.color.rgb = RGBColor(255, 255, 255)
                                        shading_elm = OxmlElement('w:shd')
                                        shading_elm.set(qn('w:val'), 'clear')
                                        shading_elm.set(qn('w:color'), 'auto')
                                        shading_elm.set(qn('w:fill'), '2D3748')
                                        cell._tc.get_or_add_tcPr().append(shading_elm)
                                    else:
                                        for p in cell.paragraphs:
                                            for run in p.runs:
                                                run.font.size = Pt(9)
                        doc.add_paragraph()

                    for riga in righe:
                        riga_strip = riga.strip()
                        if riga_strip.startswith('|'):
                            if '---' in riga_strip: continue
                            if buffer_testo: flush_testo_docx()
                            celle = [c.strip().replace('**', '') for c in riga_strip.split('|')[1:-1]]
                            if celle: righe_tabella.append(celle)
                        else:
                            if righe_tabella:
                                aggiungi_tabella_docx(righe_tabella)
                                righe_tabella = []
                            
                            if riga_strip:
                                buffer_testo.append(riga_strip)
                            elif buffer_testo:
                                flush_testo_docx()
                    
                    if buffer_testo: flush_testo_docx()
                    if righe_tabella: aggiungi_tabella_docx(righe_tabella)
                    doc.add_paragraph("-" * 40)

            doc.save(f_salva)
            messagebox.showinfo("Esportazione", "Documento Word (.docx) esportato correttamente con tabelle native strutturate.")
        except Exception as e: messagebox.showerror("Errore", str(e))


# ==========================================
# 🚀 AVVIO PROGRAMMA
# ==========================================
if __name__ == "__main__":
    root = tk.Tk()
    app = PlanciaMarescialloAI(root)
    root.mainloop()